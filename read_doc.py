import os
from docx import Document

print(f"Current directory: {os.getcwd()}")
print(f"File exists: {os.path.exists('AP POR PARSING MAP.docx')}")
print(f"File size: {os.path.getsize('AP POR PARSING MAP.docx')} bytes")

try:
    doc = Document('AP POR PARSING MAP.docx')
    print(f"Document loaded successfully. Number of paragraphs: {len(doc.paragraphs)}")
    
    print("\nPOR PARSING MAP CONTENT:")
    print("=" * 50)
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"Paragraph {i+1}: {text}")
    
    print(f"\nNumber of tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i+1} ({len(table.rows)} rows):")
        for j, row in enumerate(table.rows):
            row_text = [cell.text.strip() for cell in row.cells]
            print(f"  Row {j+1}: {' | '.join(row_text)}")
            
except Exception as e:
    print(f"Error reading document: {e}")
    import traceback
    traceback.print_exc() 